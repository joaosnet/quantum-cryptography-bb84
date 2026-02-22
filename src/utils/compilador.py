# Script de automação para injetar Markdown no DOCX do COBENGE e converter para PDF
# Dependências: pip install python-docx markdown PyMuPDF rich docx2pdf

import os
import re
import shutil
import subprocess

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from rich.console import Console
from rich.panel import Panel

# Tentativa de importar o conversor de PDF para PNG (Imagens dos Gráficos)
try:
    import fitz  # PyMuPDF

    PYMUPDF_INSTALADO = True
except ImportError:
    PYMUPDF_INSTALADO = False

# Tentativa de importar o conversor Word
try:
    from docx2pdf import convert

    DOCX2PDF_INSTALADO = True
except ImportError:
    DOCX2PDF_INSTALADO = False

console = Console()


def limpar_documento(doc):
    for paragraph in doc.paragraphs:
        p = paragraph._element
        p.getparent().remove(p)
    for table in doc.tables:
        tbl = table._element
        tbl.getparent().remove(tbl)


def adicionar_texto_markdown(paragraph, texto):
    partes = re.split(r"(\*\*.*?\*\*|\*[^*]+\*)", texto)
    for parte in partes:
        if parte.startswith("**") and parte.endswith("**"):
            run = paragraph.add_run(parte[2:-2])
            run.bold = True
        elif parte.startswith("*") and parte.endswith("*"):
            run = paragraph.add_run(parte[1:-1])
            run.italic = True
        else:
            run = paragraph.add_run(parte)

        run.font.name = "Arial"
        run.font.size = Pt(12)


def aplicar_estilo_paragrafo_cobenge(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.first_line_indent = Inches(0.49)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0


def converter_pdf_para_png(caminho_pdf):
    if not PYMUPDF_INSTALADO:
        console.print(
            "[bold red]Aviso:[/bold red] 'PyMuPDF' ausente. Rode: pip install PyMuPDF"
        )
        return caminho_pdf

    caminho_png = caminho_pdf.replace(".pdf", ".png")
    if not os.path.exists(caminho_png):
        console.print(
            f"[dim]Convertendo gráfico para PNG: {os.path.basename(caminho_pdf)}...[/dim]"
        )
        try:
            doc = fitz.open(caminho_pdf)
            page = doc.load_page(0)
            pix = page.get_pixmap(dpi=300)
            pix.save(caminho_png)
            doc.close()
        except (OSError, ValueError) as e:
            console.print(
                f"[bold red]Falha na conversão de {caminho_pdf}:[/bold red] {e}"
            )
            return caminho_pdf
    return caminho_png


def inserir_imagem(doc, caminho_imagem, legenda):
    if caminho_imagem.lower().endswith(".pdf"):
        caminho_imagem = converter_pdf_para_png(caminho_imagem)

    try:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(caminho_imagem, width=Inches(5.5))
    except (OSError, ValueError) as e:
        console.print(
            f"[bold red]Erro ao inserir a imagem {caminho_imagem}:[/bold red] {e}"
        )
        p = doc.add_paragraph(f"[FALTA A IMAGEM: {os.path.basename(caminho_imagem)}]")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p_legenda = doc.add_paragraph(legenda)
    p_legenda.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_legenda.paragraph_format.first_line_indent = Inches(0)
    for run in p_legenda.runs:
        run.font.name = "Arial"
        run.font.size = Pt(10)


def gerar_pdf_libreoffice(input_docx, output_pdf):
    """Motor de fallback usando o LibreOffice Headless (Silencioso)"""
    soffice = shutil.which("soffice")
    if not soffice:
        # Busca caminhos padrão do LibreOffice no Windows
        paths = [
            r"C:\Program Files\LibreOffice\program\soffice.exe",
            r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
        ]
        for p in paths:
            if os.path.exists(p):
                soffice = p
                break

    if soffice:
        out_dir = os.path.dirname(output_pdf)
        try:
            # Comando mágico de terminal do LibreOffice para PDF
            subprocess.run(
                [
                    soffice,
                    "--headless",
                    "--convert-to",
                    "pdf",
                    input_docx,
                    "--outdir",
                    out_dir,
                ],
                check=True,
                stdout=subprocess.DEVNULL,
                stderr=subprocess.DEVNULL,
            )
        except (subprocess.CalledProcessError, FileNotFoundError) as e:
            console.print(f"[bold red]Erro no LibreOffice:[/bold red] {e}")
            return False
        else:
            return True
    return False


def compilar_artigo(template_path, markdown_path, output_docx, output_pdf, root_dir):
    console.print("[cyan]Lendo e limpando template oficial...[/cyan]")
    doc = Document(template_path)
    limpar_documento(doc)

    console.print(
        f"[cyan]Processando o artigo: {os.path.basename(markdown_path)}...[/cyan]"
    )
    with open(markdown_path, "r", encoding="utf-8") as f:
        linhas = f.readlines()

    primeira_linha = True
    for linha in linhas:
        linha = linha.strip()
        if not linha:
            continue

        if primeira_linha and not linha.startswith("#") and not linha.startswith("!["):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(linha.upper())
            run.font.name = "Arial"
            run.font.size = Pt(14)
            run.bold = True
            p.paragraph_format.space_after = Pt(36)
            primeira_linha = False
            continue
        primeira_linha = False

        if linha.startswith("## "):
            texto_secao = linha.replace("## ", "").upper()
            p = doc.add_paragraph()
            run = p.add_run(texto_secao)
            run.font.name = "Arial"
            run.font.size = Pt(12)
            run.bold = True
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(12)
            if "REFER" in texto_secao:
                p.paragraph_format.page_break_before = True

        elif "![" in linha and "]]" in linha.replace(r"\]", "]"):
            linha_limpa = linha.replace(r"\[", "[").replace(r"\]", "]")
            conteudo = linha_limpa[linha_limpa.find("[[") + 2 : linha_limpa.find("]]")]
            if "|" in conteudo:
                caminho, legenda = conteudo.split("|", 1)
            else:
                caminho, legenda = conteudo, "Figura"
            caminho_completo = os.path.normpath(os.path.join(root_dir, caminho))
            inserir_imagem(doc, caminho_completo, legenda)

        else:
            p = doc.add_paragraph()
            aplicar_estilo_paragrafo_cobenge(p)
            adicionar_texto_markdown(p, linha)

    console.print("[yellow]Salvando arquivo DOCX...[/yellow]")
    doc.save(output_docx)
    console.print(
        f"[bold green]✔ DOCX gerado: {os.path.basename(output_docx)}[/bold green]"
    )

    # Geração Inteligente do PDF (Word -> LibreOffice)
    console.print(
        "[yellow]Invocando motor PDF para gerar o formato Acadêmico...[/yellow]"
    )
    sucesso_pdf = False

    if DOCX2PDF_INSTALADO:
        try:
            convert(output_docx, output_pdf)
            sucesso_pdf = True
        except Exception as exc:  # noqa: BLE001
            console.print(f"[bold red]Erro ao converter DOCX para PDF com docx2pdf:[/bold red] {exc}")

    if not sucesso_pdf:
        console.print(
            "[dim]Motor do MS Word indisponível. Acionando motor do LibreOffice (Headless)...[/dim]"
        )
        sucesso_pdf = gerar_pdf_libreoffice(output_docx, output_pdf)

    if sucesso_pdf:
        console.print(
            f"[bold green]✔ PDF gerado com sucesso: {os.path.basename(output_pdf)}[/bold green]"
        )
    else:
        console.print(
            "[bold red]✖ Falha na automação do Word e do LibreOffice.[/bold red]"
        )
        console.print(
            "[dim]-> Abra o DOCX gerado e salve manualmente como PDF para submeter ao COBENGE.[/dim]"
        )


if __name__ == "__main__":
    console.print(
        Panel.fit(
            "[bold magenta]Motor de Compilação COBENGE[/bold magenta]\n[dim]Suporte Híbrido: MS Word e LibreOffice Headless[/dim]",
            border_style="cyan",
        )
    )

    ROOT_DIR = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", ".."))

    template_docx = os.path.join(
        ROOT_DIR, "templates", "COBENGE-2025-Template-STe-SP.docx"
    )
    artigo_md = os.path.join(ROOT_DIR, "docs", "artigo.md")

    saida_docx = os.path.join(ROOT_DIR, "Projeto_Final_BB84_COBENGE.docx")
    saida_pdf = os.path.join(ROOT_DIR, "Projeto_Final_BB84_COBENGE.pdf")

    compilar_artigo(template_docx, artigo_md, saida_docx, saida_pdf, ROOT_DIR)
