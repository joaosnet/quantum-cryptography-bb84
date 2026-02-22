# Script de automação para injetar Markdown no Template Oficial DOCX do COBENGE
# Dependências: pip install python-docx markdown PyMuPDF

import os
import sys

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from rich.console import Console
from rich.panel import Panel

# Tentativa de importar o conversor de PDF
try:
    import fitz  # PyMuPDF

    PYMUPDF_INSTALADO = True
except ImportError:
    PYMUPDF_INSTALADO = False

console = Console()


def aplicar_estilo_cobenge(paragraph):
    """Garante que o parágrafo siga as regras de recuo e fonte do COBENGE"""
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    # O COBENGE pede recuo na primeira linha de 1.25cm
    paragraph.paragraph_format.first_line_indent = Inches(0.49)  # 1.25 cm
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0  # Espaçamento simples

    # Força a fonte Arial tamanho 12
    for run in paragraph.runs:
        run.font.name = "Arial"
        run.font.size = Pt(12)


def converter_pdf_para_png(caminho_pdf):
    """Verificador: Checa se a imagem é PDF e converte para PNG usando PyMuPDF"""
    if not PYMUPDF_INSTALADO:
        console.print(
            "[bold red]Aviso Crítico:[/bold red] Biblioteca 'PyMuPDF' não instalada. Não é possível converter PDFs automaticamente. Rode: pip install PyMuPDF"
        )
        return caminho_pdf

    caminho_png = caminho_pdf.replace(".pdf", ".png")

    # Se o PNG já foi convertido em execuções anteriores, pula para economizar CPU
    if not os.path.exists(caminho_png):
        console.print(
            f"[yellow]Verificador ativado: Convertendo PDF incompatível para PNG... ({os.path.basename(caminho_pdf)})[/yellow]"
        )
        try:
            doc = fitz.open(caminho_pdf)
            page = doc.load_page(0)  # Pega apenas a primeira página do PDF (onde está o gráfico)
            pix = page.get_pixmap(dpi=300)  # Renderiza em 300 DPI para não perder qualidade no Word
            pix.save(caminho_png)
            doc.close()
        except (RuntimeError, OSError, ValueError) as e:
            console.print(
                f"[bold red]Falha ao converter o arquivo {caminho_pdf}:[/bold red] {e}"
            )
            return caminho_pdf  # Em caso de erro catastrófico, retorna o PDF original

    return caminho_png


def inserir_imagem(doc, caminho_imagem, legenda):
    """Insere uma imagem centralizada com legenda"""

    # 1. PASSO DE VERIFICAÇÃO E CONVERSÃO
    if caminho_imagem.lower().endswith(".pdf"):
        caminho_imagem = converter_pdf_para_png(caminho_imagem)

    # 2. INSERÇÃO NO WORD
    try:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        # O python-docx vai aceitar tranquilamente porque agora é um PNG
        run.add_picture(caminho_imagem, width=Inches(5.5))
    except (OSError, ValueError) as e:
        console.print(
            f"[bold red]Erro ao inserir imagem {caminho_imagem} no DOCX:[/bold red] {e}"
        )
        p = doc.add_paragraph(f"[FALTA A IMAGEM: {os.path.basename(caminho_imagem)}]")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Legenda da imagem formatada nos padrões acadêmicos
    p_legenda = doc.add_paragraph(legenda)
    p_legenda.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_legenda.paragraph_format.first_line_indent = Inches(0)
    for run in p_legenda.runs:
        run.font.name = "Arial"
        run.font.size = Pt(10)


def compilar_artigo(template_path, markdown_path, output_path):
    if not os.path.exists(template_path):
        console.print(
            f"[bold red]✖ Template não encontrado: {template_path}[/bold red]"
        )
        sys.exit(1)

    if not os.path.exists(markdown_path):
        console.print(
            f"[bold red]✖ Arquivo Markdown não encontrado: {markdown_path}[/bold red]"
        )
        sys.exit(1)

    console.print("[cyan]Abrindo template oficial do COBENGE...[/cyan]")
    doc = Document(template_path)

    # Inicia a injeção na página seguinte (após as instruções da capa do template)
    doc.add_page_break()

    console.print(
        f"[cyan]Lendo conteúdo em Markdown: {os.path.basename(markdown_path)}[/cyan]"
    )
    with open(markdown_path, "r", encoding="utf-8") as f:
        linhas = f.readlines()

    console.print(
        "[yellow]Injetando conteúdo e aplicando formatação estrutural...[/yellow]"
    )

    for linha in linhas:
        linha = linha.strip()
        if not linha:
            continue

        # Verifica se é um cabeçalho (Seção)
        if linha.startswith("## "):
            texto_secao = linha.replace("## ", "").upper()
            p = doc.add_paragraph()
            run = p.add_run(texto_secao)
            run.font.name = "Arial"
            run.font.size = Pt(12)
            run.bold = True
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(12)

        # Verifica se é a nossa marcação de imagem customizada do Obsidian/Markdown
        elif linha.startswith("![[") and linha.endswith("]]"):
            conteudo = linha[3:-2]
            if "|" in conteudo:
                caminho, legenda = conteudo.split("|", 1)
            else:
                caminho = conteudo
                legenda = "Figura sem legenda"

            # Normaliza o caminho para rodar perfeitamente independente da pasta raiz
            caminho_completo = os.path.normpath(
                os.path.join(os.path.dirname(markdown_path), "..", caminho)
            )
            inserir_imagem(doc, caminho_completo, legenda)

        # Se for texto de parágrafo normal
        else:
            p = doc.add_paragraph(linha)
            aplicar_estilo_cobenge(p)

    console.print(f"[cyan]Salvando documento final em:[/cyan] {output_path}")
    doc.save(output_path)
    console.print(
        "[bold green]✔ Artigo compilado com sucesso para DOCX! As imagens foram tratadas.[/bold green]"
    )


if __name__ == "__main__":
    console.print(
        Panel.fit(
            "[bold magenta]Injetor de Markdown em DOCX - COBENGE[/bold magenta]\n[dim]Com Verificador Automático de PDF para PNG (PyMuPDF)[/dim]",
            border_style="cyan",
        )
    )

    # Caminhos baseados na sua File Tree
    ROOT_DIR = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", ".."))

    template_docx = os.path.join(
        ROOT_DIR, "templates", "COBENGE-2025-Template-STe-SP.docx"
    )
    artigo_md = os.path.join(ROOT_DIR, "docs", "artigo.md")
    saida_docx = os.path.join(ROOT_DIR, "Projeto_Final_BB84_COBENGE.docx")

    compilar_artigo(template_docx, artigo_md, saida_docx)
